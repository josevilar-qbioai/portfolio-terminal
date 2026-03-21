"""
Genera el paper SSRN: Escasez y Resiliencia
Formato: working paper académico, ~14 páginas
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Márgenes ──────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── Estilos base ──────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(11)

def set_heading_style(style_name, size, bold=True, color=None):
    s = doc.styles[style_name]
    s.font.name = 'Times New Roman'
    s.font.size = Pt(size)
    s.font.bold = bold
    if color:
        s.font.color.rgb = RGBColor(*color)
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after  = Pt(6)

set_heading_style('Heading 1', 13, bold=True)
set_heading_style('Heading 2', 12, bold=True)
set_heading_style('Heading 3', 11, bold=True)

def add_para(text, style='Normal', bold=False, italic=False, size=None,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6,
             color=None, indent=None):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent is not None:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    if size:  run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    return p

def add_formula(text):
    """Párrafo centrado con fórmula en negrita monoespacio."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(12)
    return p

def add_table(headers, rows, col_widths):
    """Tabla simple con cabecera gris."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Cabecera
    hdr_row = t.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr_row.cells[i]
        cell.width = Cm(w)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        # Fondo gris
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9D9D9')
        tcPr.append(shd)
    # Filas de datos
    for r_i, row_data in enumerate(rows):
        row = t.rows[r_i + 1]
        for c_i, (cell_text, w) in enumerate(zip(row_data, col_widths)):
            cell = row.cells[c_i]
            cell.width = Cm(w)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
    doc.add_paragraph()
    return t

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent  = Cm(1 + level * 0.5)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    return p

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run('─' * 72)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(180, 180, 180)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(40)
p_title.paragraph_format.space_after  = Pt(8)
r = p_title.add_run('Escasez y Resiliencia:')
r.bold = True; r.font.size = Pt(20); r.font.name = 'Times New Roman'

p_title2 = doc.add_paragraph()
p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title2.paragraph_format.space_after = Pt(4)
r2 = p_title2.add_run('Un Marco Matemático para Invertir en la Era de la')
r2.bold = True; r2.font.size = Pt(16); r2.font.name = 'Times New Roman'

p_title3 = doc.add_paragraph()
p_title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title3.paragraph_format.space_after = Pt(30)
r3 = p_title3.add_run('Autorreplicación Artificial')
r3.bold = True; r3.font.size = Pt(16); r3.font.name = 'Times New Roman'

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(8)
rs = p_sub.add_run('Scarcity and Resilience: A Mathematical Framework for Investing\nin the Age of Artificial Self-Replication')
rs.italic = True; rs.font.size = Pt(12); rs.font.name = 'Times New Roman'

add_separator()

p_auth = doc.add_paragraph()
p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_auth.paragraph_format.space_before = Pt(20)
p_auth.paragraph_format.space_after  = Pt(6)
ra = p_auth.add_run('Jose Antonio Vilar')
ra.bold = True; ra.font.size = Pt(13); ra.font.name = 'Times New Roman'

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_date.paragraph_format.space_after = Pt(4)
rd = p_date.add_run('Profesional de Relaciones con Inversores · Estudiante de Grado en Ciencia de Datos, UOC')
rd.font.size = Pt(11); rd.font.name = 'Times New Roman'

p_email = doc.add_paragraph()
p_email.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_email.paragraph_format.space_after = Pt(4)
re_ = p_email.add_run('javsprivate@gmail.com · joseantoniovilar@uoc.edu')
re_.font.size = Pt(10); re_.font.name = 'Times New Roman'; re_.italic = True

p_date2 = doc.add_paragraph()
p_date2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_date2.paragraph_format.space_after = Pt(30)
rd2 = p_date2.add_run('Marzo 2026')
rd2.font.size = Pt(11); rd2.font.name = 'Times New Roman'

add_separator()

# JEL + keywords
p_jel = doc.add_paragraph()
p_jel.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_jel.paragraph_format.space_before = Pt(16)
p_jel.paragraph_format.space_after  = Pt(4)
rj = p_jel.add_run('Clasificación JEL: G11 · G12 · O33 · Q02')
rj.font.size = Pt(10); rj.font.name = 'Times New Roman'; rj.italic = True

p_kw = doc.add_paragraph()
p_kw.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_kw.paragraph_format.space_after = Pt(4)
rk = p_kw.add_run('Palabras clave: construcción de cartera, autorreplicación logística, escasez de recursos, Bitcoin, inteligencia artificial, difusión tecnológica')
rk.font.size = Pt(10); rk.font.name = 'Times New Roman'; rk.italic = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('Abstract', level=1)

add_para(
    'This paper proposes a mathematical framework for constructing a long-term investment '
    'portfolio at the intersection of accelerating artificial intelligence and physical resource '
    'scarcity. We introduce the Scarcity and Resilience model, V(t) = Capital × (1+r)^t × Φ_L(t), '
    'where r is the AI ecosystem growth rate (calibrated via the Philadelphia Semiconductor Index, SOX) '
    'and Φ_L(t) = 1 + K/(1+e^(−γ(t−t₀))) is a logistic autoreplication factor governed by three '
    'independently calibrable parameters: K (maximum replication multiplier), γ (adoption velocity, '
    'calibrated via major LLM release cadence), and t₀ (inflection year, detectable via the '
    'SOX/Copper Divergence signal). Following the tradition of Bass (1969) diffusion models, '
    'this logistic specification provides an upper bound to autoreplication (Φ_L → 1+K as t → ∞), '
    'separates compound growth from technology adoption dynamics, and yields three independently '
    'falsifiable parameters. The portfolio is structured across five thesis pillars: Digital Scarcity '
    '(Bitcoin), AI Autoreplication (semiconductors, robotics), Energy/Grid (uranium, smart grid '
    'utilities), Physical Scarcity (copper, gold), and Resilience (global equity funds). We further '
    'introduce the SOX/Copper Correlation as a paradigm-shift detector whose Divergence state '
    '(SOX rising, Copper falling) marks the empirical t₀, triggering portfolio rotation from '
    'physical to digital scarcity assets. The framework is falsifiable: if AI adoption stalls, '
    'Φ_L → 1 and the model reduces to standard compound growth.',
    italic=True
)

doc.add_paragraph()

doc.add_heading('Resumen', level=1)
add_para(
    'Este artículo propone un marco matemático para la construcción de una cartera de inversión '
    'a largo plazo en la intersección de la aceleración de la inteligencia artificial y la escasez '
    'de recursos físicos. Introducimos el modelo Escasez y Resiliencia, V(t) = Capital × (1+r)^t × Φ_L(t), '
    'donde r es la tasa de crecimiento del ecosistema de IA (calibrada mediante el índice SOX) y '
    'Φ_L(t) = 1 + K/(1+e^(−γ(t−t₀))) es un factor de autorreplicación logístico gobernado por '
    'tres parámetros calibrables de forma independiente: K (multiplicador máximo de autorreplicación), '
    'γ (velocidad de adopción, calibrada desde la cadencia de releases LLM) y t₀ (año de inflexión, '
    'detectable empíricamente mediante la señal de Divergencia SOX/Cobre). Siguiendo la tradición '
    'de los modelos de difusión de Bass (1969), la especificación logística establece un techo a '
    'la autorreplicación (Φ_L → 1+K cuando t → ∞), separa el crecimiento compuesto de la dinámica '
    'de adopción tecnológica y proporciona tres parámetros independientemente falsificables. '
    'La cartera se estructura en cinco pilares: Escasez Digital (Bitcoin), Autorreplicación IA '
    '(semiconductores, robótica), Energía/Grid (uranio, utilities de smart grid), Escasez Física '
    '(cobre, oro) y Resiliencia (fondos indexados globales). Presentamos además la Correlación '
    'SOX/Cobre como detector de cambio de paradigma, cuyo estado de Divergencia marca '
    'empíricamente t₀ y activa la rotación táctica de la cartera. El modelo es falsificable: '
    'si la adopción de IA se estanca, Φ_L → 1 y el modelo converge al crecimiento compuesto estándar.',
    italic=True
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Introducción', level=1)

doc.add_heading('1.1 Origen: la intuición precede al modelo', level=2)

add_para(
    'Este paper surge de un proceso inverso al habitual en la investigación académica financiera. '
    'El punto de partida no fue una hipótesis teórica que se buscó contrastar empíricamente, '
    'sino una cartera de inversión real, construida a lo largo de varios años a partir de una '
    'intuición sobre la estructura del mercado durante la década de los años veinte del siglo XXI. '
    'El modelo matemático presentado en las secciones siguientes fue desarrollado con posterioridad, '
    'con el objetivo de formalizar esa intuición, identificar sus supuestos implícitos, calibrar '
    'sus parámetros y determinar las condiciones bajo las cuales la tesis podría ser falsificada.'
)

add_para(
    'Esta secuencia — cartera primero, modelo después — no es infrecuente en finanzas. Benjamin '
    'Graham formalizó los principios del value investing décadas después de aplicarlos con éxito '
    'como gestor. Harry Markowitz desarrolló la teoría moderna de carteras (1952) a partir de '
    'observaciones prácticas sobre la relación entre diversificación y riesgo. La presente '
    'contribución es más modesta en sus pretensiones, pero sigue la misma lógica: el modelo es '
    'un instrumento de disciplina y comunicación de una convicción inversora, no una predicción.'
)

add_para(
    'La declaración de conflicto de interés relevante es la siguiente: el autor mantiene '
    'posiciones en todos los activos mencionados en el paper. La cartera fue construida antes '
    'de que existiera el modelo matemático que aquí se presenta. El lector debe interpretar '
    'el framework propuesto como la racionalización post-hoc de una intuición, no como un '
    'sistema de selección de activos derivado deductivamente de axiomas independientes.'
)

doc.add_heading('1.2 La intuición central', level=2)

add_para(
    'La intuición que motivó la construcción de la cartera puede enunciarse en dos proposiciones:'
)

add_bullet(
    'Proposición 1 — Ventana de complementariedad: durante el período de despliegue masivo de '
    'la inteligencia artificial (aproximadamente 2020-2035), la demanda de activos de escasez '
    'física (cobre, uranio, metales preciosos) y de activos de escasez digital (Bitcoin) crece '
    'simultáneamente, porque la IA necesita infraestructura física para entrenarse y ejecutarse, '
    'y al mismo tiempo genera devaluación monetaria a través de la aceleración de la productividad.'
)

add_bullet(
    'Proposición 2 — Cambio de paradigma detectable: existe un indicador de mercado observable '
    'que señala el fin de esta ventana de complementariedad: la Divergencia SOX/Cobre, definida '
    'como el momento en que el índice de semiconductores sigue apreciándose mientras el precio '
    'del cobre inicia una tendencia bajista estructural. Dicha divergencia indicaría que la IA '
    'ha comenzado a miniaturizar o sustituir su demanda de materiales físicos.'
)

add_para(
    'El modelo matemático presentado en este paper tiene como único propósito formalizar estas '
    'dos proposiciones de manera que sean calibrables, comparables entre escenarios y, '
    'crucialmente, falsificables.'
)

doc.add_heading('1.3 La tensión física-digital y la pregunta central', level=2)

add_para(
    'La convergencia entre la aceleración exponencial de la inteligencia artificial y la '
    'escasez estructural de recursos físicos plantea una de las tensiones económicas más '
    'relevantes de la próxima década. Por un lado, el despliegue masivo de centros de datos, '
    'la fabricación de chips avanzados y la electrificación global de la industria generan '
    'una demanda sin precedentes de materiales como el cobre, el uranio y los metales preciosos '
    '(IEA, 2021). Por otro, la propia IA está acelerando los ciclos de innovación en materiales, '
    'miniaturización y eficiencia energética (Brynjolfsson y McAfee, 2014), lo que podría '
    'reducir esa demanda física en el futuro.'
)

add_para(
    '¿Cuándo la IA dejará de necesitar "hierro"? Esta es la pregunta central que articula la '
    'tesis. La hipótesis de trabajo es que existe una ventana temporal, de duración incierta '
    'pero finita, durante la cual ambas categorías de activos se aprecian simultáneamente. '
    'Identificar el inicio del fin de esa ventana — con suficiente antelación para rotar '
    'la cartera — es el objetivo estratégico del framework.'
)

add_para(
    'El paper se estructura del siguiente modo: la Sección 2 presenta el modelo matemático '
    'Escasez y Resiliencia con su factor de autorreplicación logístico. La Sección 3 '
    'describe los cinco pilares de la cartera. La Sección 4 introduce la Correlación SOX/Cobre '
    'como detector de cambio de paradigma. La Sección 5 desarrolla la calibración de parámetros. '
    'La Sección 6 proyecta los tres escenarios. La Sección 7 analiza los factores de riesgo '
    'y limitaciones. La Sección 8 concluye.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. EL MODELO MATEMÁTICO
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('2. El Modelo Matemático: Escasez y Resiliencia', level=1)

doc.add_heading('2.1 Formulación general', level=2)

add_para(
    'El modelo propone que el valor de la cartera en el momento t viene dado por:'
)

add_formula('V(t)  =  Capital  ×  (1 + r_IA)^t  ×  Φ(t)')

add_para(
    'donde V(t) es el valor proyectado de la cartera en el año t, Capital es el valor '
    'actual de la cartera, r_IA es la tasa de crecimiento anual del ecosistema de inteligencia '
    'artificial (calibrada desde el índice SOX), y Φ(t) es el factor de autorreplicación, '
    'que captura la aceleración no lineal producida por la automejora de sistemas de IA.'
)

add_para(
    'El componente (1+r_IA)^t es el factor de crecimiento compuesto estándar, análogo al '
    'utilizado en modelos de valoración de activos de crecimiento. Cuando Φ(t) = 1, '
    'el modelo converge al crecimiento compuesto estándar — esta es la hipótesis nula '
    'que el modelo debe superar empíricamente para ser útil.'
)

doc.add_heading('2.2 El factor de autorreplicación Φ_L(t): curva logística', level=2)

add_para(
    'La elección de una curva logística (S-curve) para Φ(t) responde a tres requisitos '
    'conceptuales que cualquier modelo de adopción tecnológica debe satisfacer: (i) un período '
    'inicial de arranque lento, (ii) una fase de aceleración máxima en torno a un año de '
    'inflexión t₀, y (iii) una saturación final acotada por factores termodinámicos, económicos '
    'y regulatorios. Esta estructura sigue la tradición de los modelos de difusión tecnológica '
    'de Bass (1969) y Rogers (2003), ampliamente validados en la literatura de innovación. '
    'El modelo propone para Φ_L(t):'
)

add_formula('Φ_L(t)  =  1  +  K / (1 + e^(−γ·(t − t₀)))')

add_para(
    'donde K es el multiplicador máximo de autorreplicación (cuánto amplifica la IA '
    'el crecimiento base en su estado maduro), γ es la velocidad de adopción (calibrada '
    'desde la cadencia de releases LLM), y t₀ es el año de inflexión '
    '(el punto de máxima aceleración del paradigma). Los tres parámetros tienen '
    'interpretación económica directa e independiente:'
)

add_bullet(
    'K acota el crecimiento: cuando t → ∞, Φ_L → 1 + K. La autorreplicación no puede '
    'crecer indefinidamente — tiene un límite termodinámico, económico y regulatorio. '
    'K se calibra como ratio de expansión del mercado de IA respecto al mercado actual.'
)
add_bullet(
    'γ determina la velocidad: captura cuán rápidamente el ecosistema alcanza la máxima '
    'aceleración. Se calibra desde la cadencia de releases de modelos LLM de frontera: '
    'un release cada 6 meses implica γ ≈ 0.5; cada 3 meses implica γ ≈ 0.9.'
)
add_bullet(
    'El modelo es falsificable: t₀ es observable a posteriori mediante la señal '
    'SOX/Cobre (Divergencia sostenida > 3 meses). Si la adopción se estanca, '
    'Φ_L(t) → 1 y el modelo converge al crecimiento compuesto estándar.'
)

doc.add_heading('2.3 Función de valor por categoría', level=2)

add_para(
    'Dado que cada pilar de la cartera tiene una exposición diferencial al factor de '
    'autorreplicación, el modelo se aplica también por categoría de activo, con parámetros '
    'K_cat y γ_cat específicos que reflejan su distinta velocidad y magnitud de adopción '
    'del paradigma de IA:'
)

add_formula('V_cat(t)  =  V_cat(0)  ×  (1 + r_cat)^t  ×  Φ_L(t; K_cat, γ_cat, t₀_cat)')

add_para(
    'La Tabla 1 resume los parámetros estimados por categoría. Los valores de K '
    'representan el multiplicador máximo de autorreplicación atribuible a cada pilar. '
    'Los valores de γ_base y γ_opt corresponden a los escenarios Base y Óptimo respectivamente.'
)
doc.add_paragraph()

add_table(
    headers=['Categoría', 'r_cat', 'K Base', 'K Óptimo', 'γ_base', 'Pilar de la tesis'],
    rows=[
        ['Bitcoin / Crypto',      '15%', '2.0', '6.0', '0.5', 'Escasez monetaria digital'],
        ['IA / Robótica',         '25%', '4.0', '9.0', '1.0', 'Autorreplicación directa'],
        ['Energía / Nuclear+Grid','20%', '2.5', '5.0', '0.7', 'Combustible de la tesis'],
        ['Metales (Cu, Au)',      '10%', '1.5', '3.0', '0.4', 'Escasez física irreducible'],
        ['Renta Variable Global', '8%',  '0.5', '1.5', '0.3', 'Resiliencia pasiva'],
    ],
    col_widths=[4.2, 1.4, 1.4, 1.7, 1.5, 4.8]
)

add_para('Tabla 1. Parámetros del modelo logístico por categoría de activo. '
         'K = multiplicador máximo de autorreplicación. γ = velocidad de adopción (calibrada desde cadencia LLM).',
         italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)

# ══════════════════════════════════════════════════════════════════════════════
# 3. LOS CINCO PILARES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Los Cinco Pilares de la Cartera', level=1)

doc.add_heading('3.1 Escasez Digital: Bitcoin', level=2)
add_para(
    'Bitcoin posee una oferta máxima de 21 millones de unidades, programada matemáticamente '
    'en su protocolo y no sujeta a discrecionalidad política ni tecnológica. En un entorno '
    'de creciente emisión monetaria fiat y de adopción institucional creciente, Bitcoin cumple '
    'la función de reserva de valor escasa de la misma forma que el oro en épocas anteriores, '
    'pero con el añadido de ser nativo del entorno digital donde la IA opera. '
    'La hipótesis adicional es que, en un escenario de agentes de IA operando de forma '
    'autónoma, Bitcoin se posiciona como la "moneda de las máquinas": un medio de intercambio '
    'que los sistemas de IA pueden utilizar sin intermediación humana ni riesgo de censura.'
)

doc.add_heading('3.2 Autorreplicación IA: Semiconductores y Robótica', level=2)
add_para(
    'Los activos de semiconductores, inteligencia artificial y robótica son la exposición '
    'directa al factor de autorreplicación Φ_L(t) del modelo, y en particular al parámetro K '
    '— el multiplicador máximo de autorreplicación — que toma sus valores más elevados en '
    'esta categoría (K_base=4.0, K_opt=9.0, según Tabla 1). Las empresas que fabrican los chips '
    'sobre los que se entrena y ejecuta la IA, y las que desarrollan los robots que pueden '
    'eventualmente fabricar otros robots, son beneficiarias directas de la aceleración '
    'del ecosistema, cuya velocidad de adopción γ se calibra precisamente desde la cadencia '
    'de releases LLM. El índice SOX (Philadelphia Semiconductor Index) actúa como proxy de '
    'mercado del ritmo de inversión en capex de IA, siendo a su vez un indicador adelantado '
    'de la demanda energética futura.'
)

doc.add_heading('3.3 Energía / Grid: Uranio y Utilities', level=2)
add_para(
    'El despliegue masivo de centros de datos de IA requiere energía continua, densa y '
    'predecible. La energía nuclear de fisión — especialmente los Small Modular Reactors '
    '(SMRs) y los reactores Gen-IV — se posiciona como la fuente más viable para satisfacer '
    'este requisito, dado que el gas natural introduce volatilidad de precio y las energías '
    'renovables no proporcionan la continuidad de suministro que requiere el cómputo de alto '
    'rendimiento. El uranio (proxy: Cameco, CCJ) es el input escaso de esta cadena. '
    'Las utilities estadounidenses (ETF XLU) se benefician de los Power Purchase Agreements '
    '(PPAs) firmados con hyperscalers como Google, Microsoft y Amazon.'
)
add_para(
    'La incorporación de infraestructura de smart grid (redes eléctricas inteligentes) '
    'representa la capa de distribución de esta energía: sin redes capaces de gestionar '
    'flujos bidireccionales y demanda variable, la expansión del ecosistema de IA encuentra '
    'un cuello de botella físico.'
)

doc.add_heading('3.4 Escasez Física: Cobre y Oro', level=2)
add_para(
    'El cobre es el material de la electrificación: cada robot, cada vehículo eléctrico, '
    'cada centro de datos y cada aerogenerador requiere cantidades significativas de cobre. '
    'La demanda estructural de cobre es, por tanto, una función directa del despliegue '
    'físico de la IA. El oro mantiene su función histórica de refugio de valor ante la '
    'devaluación monetaria y actúa como cobertura ante escenarios de inestabilidad '
    'macroeconómica que podrían interrumpir la tesis principal.'
)

doc.add_heading('3.5 Resiliencia: Renta Variable Global', level=2)
add_para(
    'La inclusión de fondos indexados globales (MSCI World, S&P 500) cumple una función '
    'de resiliencia: si la tesis principal resulta incorrecta en sus parámetros temporales, '
    'la cartera mantiene exposición al crecimiento económico general. Esta categoría actúa '
    'como amortiguador ante escenarios de desaceleración del ecosistema de IA que no '
    'impliquen una recesión generalizada.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. LA CORRELACIÓN SOX / COBRE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('4. La Correlación SOX/Cobre: Detector de Cambio de Paradigma', level=1)

add_para(
    'El indicador más relevante para la gestión dinámica de la cartera es la correlación '
    'entre el rendimiento mensual del índice SOX (semiconductores) y el precio del cobre '
    '(HG1!, contrato de futuros de cobre de primer vencimiento). Esta correlación actúa '
    'como detector de cambio de paradigma tecnológico.'
)

add_para(
    'La lógica es la siguiente: durante la "Era de Construcción" — el período actual, en '
    'el que la IA requiere masiva infraestructura física para entrenar y desplegar modelos '
    '— ambos activos se aprecian simultáneamente, ya que la demanda de chips y la demanda '
    'de cobre son complementarias. Sin embargo, existe un punto de inflexión potencial: '
    'cuando la IA alcance suficiente capacidad de miniaturización y diseño de materiales '
    'como para sustituir o reducir significativamente el uso de cobre (y otros metales '
    'críticos), el SOX puede continuar su apreciación mientras el cobre inicia un ciclo '
    'bajista. Este es el momento de rotar fuera de las posiciones en metales y '
    'concentrar el capital en los activos de autorreplicación.'
)

doc.add_paragraph()
add_table(
    headers=['Estado', 'Condición', 'Interpretación', 'Acción estratégica'],
    rows=[
        ['Era de Construcción', 'SOX ↑  y  Cobre ↑',
         'IA requiere materia física', 'Mantener todos los pilares'],
        ['Divergencia · Alerta', 'SOX ↑  /  Cobre ↓',
         'IA miniaturiza o sustituye', 'Rotar metales → IA pura'],
        ['Escasez Física', 'SOX ↓  /  Cobre ↑',
         'Demanda física sin capex IA', 'Mantener metales, esperar SOX'],
        ['Contracción', 'SOX ↓  y  Cobre ↓',
         'Entorno adverso generalizado', 'Modo defensivo — reducir riesgo'],
    ],
    col_widths=[3.5, 3.2, 4.3, 4.0]
)

add_para('Tabla 2. Estados de la Correlación SOX/Cobre y acciones estratégicas.',
         italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)

add_para(
    'La Divergencia (SOX sube, Cobre baja) es la señal más importante: es el primer '
    'indicador observable de que el paradigma está cambiando y de que el valor se '
    'está transfiriendo desde los activos físicos hacia los activos de información. '
    'Históricamente, estos cambios de paradigma se producen de forma gradual pero '
    'acelerada una vez que se confirman durante dos o tres meses consecutivos.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. CALIBRACIÓN DE PARÁMETROS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Calibración de los Parámetros r_IA, K, γ y t₀', level=1)

add_para(
    'Una de las principales mejoras del modelo logístico revisado respecto a la versión '
    'original es que sus cuatro parámetros son calibrables de forma independiente, '
    'cada uno a partir de una fuente de datos distinta y objetiva.'
)

doc.add_heading('5.1 Calibración de r_IA', level=2)
add_para(
    'El parámetro r_IA es el más directamente calibrable a partir de datos de mercado. '
    'El rendimiento anualizado del índice SOX proporciona una estimación de mercado del '
    'ritmo al que los inversores están descontando el crecimiento del ecosistema de IA. '
    'La regla de calibración propuesta es:'
)
add_formula('r_IA  =  Rendimiento_SOX_12M  /  100')
add_para(
    'Por ejemplo, si el SOX ha subido un 22% en los últimos 12 meses, r_IA = 0.22. '
    'Esta calibración se revisa trimestralmente. Como sanidad check, puede utilizarse '
    'también el rendimiento del NASDAQ Composite, que tiende a ser más estable que el SOX.'
)

doc.add_heading('5.2 Calibración de γ — velocidad de adopción', level=2)
add_para(
    'El parámetro γ es el gran avance del modelo revisado: por primera vez, la velocidad '
    'de autorreplicación puede calibrarse de forma completamente cuantitativa y objetiva, '
    'sin juicios cualitativos. La fuente de calibración es la cadencia de releases de '
    'Grandes Modelos de Lenguaje (LLMs) de capacidad fundacional superior.'
)
add_para(
    'La lógica es la siguiente: γ mide cuánto tarda el ecosistema de IA en completar '
    'un ciclo de automejoría. Cada release LLM mayor representa un ciclo completo '
    '(entrenamiento → evaluación → mejora de arquitectura). La regla de calibración es:'
)
add_formula('γ  =  12  /  intervalo_mediano_entre_releases_LLM  (meses)')
add_para(
    'Con datos hasta marzo de 2026: el intervalo mediano entre releases mayores es '
    'aproximadamente 4-5 meses, lo que produce γ ∈ [2.4, 3.0] en el escenario acelerado '
    'y γ ≈ 0.7-1.0 en el escenario base (considerando solo los releases tier-1 de cada lab). '
    'La Tabla 3 resume la escala de calibración de γ junto con los proxies de mercado '
    'correspondientes.'
)

doc.add_paragraph()
add_table(
    headers=['γ', 'Intervalo LLM', 'Condición observable', 'Proxy de mercado'],
    rows=[
        ['0.5', '~24 meses', 'Adopción lenta · releases espaciados',
         'SOX ≈ NASDAQ'],
        ['0.9', '~13 meses', 'Cadencia histórica media (2020-2024)',
         'SOX ligeramente > NASDAQ'],
        ['1.5', '~8 meses',  'Aceleración · múltiples labs compitiendo',
         'SOX > NASDAQ trailing 3M'],
        ['2.5+', '< 5 meses', 'Autorreplicación activa · releases trimestrales',
         'SOX > NASDAQ + divergencia NVDA'],
    ],
    col_widths=[1.2, 2.5, 5.5, 4.8]
)
add_para('Tabla 3. Escala de calibración del parámetro γ a partir de la cadencia de releases LLM.',
         italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)

doc.add_heading('5.3 Calibración de K — multiplicador máximo', level=2)
add_para(
    'El parámetro K es el único que no puede calibrarse de forma puramente cuantitativa: '
    'representa el multiplicador máximo que la autorreplicación de IA puede aportar al '
    'crecimiento base una vez que el paradigma está maduro. K es, en esencia, una '
    'hipótesis sobre el estado final del ecosistema. La calibración propuesta relaciona '
    'K con el potencial de expansión total del mercado de IA respecto al mercado actual:'
)
add_formula('K  ≈  (TAM_IA_maduro  /  TAM_IA_actual)  −  1')
add_para(
    'Con estimaciones de mercado actuales (TAM IA ~$200B en 2025, proyecciones '
    '$15-20T en escenario de madurez), K ∈ [3, 9] dependiendo del escenario. '
    'K tiene una interpretación económica directa (ratio de expansión de mercado) '
    'que lo hace falsificable: si el TAM de IA no crece conforme a las proyecciones, '
    'K debe reducirse en la siguiente revisión trimestral.'
)

doc.add_heading('5.4 Calibración de t₀ — año de inflexión', level=2)
add_para(
    'El parámetro t₀ es el más estratégicamente relevante: es el año en que el ecosistema '
    'alcanza su máxima tasa de aceleración y, por tanto, el momento en que la señal '
    'SOX/Cobre debería detectar la Divergencia. La calibración se realiza a partir de '
    'tres proxies convergentes:'
)
add_bullet(
    'Proyecciones de AGI / automatización de conocimiento: la mayoría de los labs '
    'de IA proyectan capacidades de nivel experto en tareas cognitivas complejas '
    'para el período 2027-2030, lo que sugiere t₀ ∈ [3, 6] años desde 2026.'
)
add_bullet(
    'Curvas de coste de inferencia: al ritmo actual de caída (10× cada 12 meses), '
    'el coste de inferencia alcanzará paridad con el trabajo humano en conocimiento '
    'aproximadamente en 2028-2029, otro proxy de t₀.'
)
add_bullet(
    'Señal de mercado SOX/Cobre: t₀ puede estimarse retrospectivamente como el año '
    'en que la Divergencia se mantiene durante más de 3 meses consecutivos, '
    'proporcionando una estimación ex-post que actualiza la calibración prospectiva.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. ESCENARIOS Y PROYECCIONES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Escenarios y Proyecciones', level=1)

add_para(
    'A partir de la formulación del modelo revisado y los parámetros de calibración descritos, '
    'se presentan tres escenarios de proyección para un horizonte de diez años. '
    'Los escenarios se diferencian en los valores de r, K, γ y t₀ asignados, reflejando '
    'distintas hipótesis sobre la velocidad de adopción (γ) y el multiplicador máximo de '
    'autorreplicación (K) del ecosistema de IA. El año de inflexión t₀ = 5 (2031) es común '
    'a todos los escenarios, consistente con las proyecciones de capacidades de nivel experto '
    'descritas en §5.4.'
)

add_table(
    headers=['Escenario', 'r_IA', 'K', 'γ', 'λ_eff', 'Factor ×2a', 'Factor ×4a', 'Factor ×10a'],
    rows=[
        ['Base',      '20%', '2.0', '0.50', '0.125', '2.0×', '3.6×',  '17.6×'],
        ['Acelerado', '25%', '4.0', '0.90', '0.300', '2.0×', '5.3×',  '46.2×'],
        ['Óptimo',    '30%', '6.0', '1.50', '0.562', '1.8×', '6.0×',  '96.5×'],
    ],
    col_widths=[2.4, 1.2, 1.1, 1.2, 1.4, 1.8, 1.8, 2.1]
)

add_para('Tabla 4. Parámetros del modelo logístico y factores de multiplicación por escenario '
         '(t₀=5, horizonte=10 años). λ_eff = γ·K/(4+2K). Factores calculados con capital=100, t=2, 4, 10.',
         italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)

add_para(
    'Es importante subrayar que los factores de multiplicación del escenario Óptimo '
    'son el resultado matemático de la combinación de una tasa de crecimiento elevada '
    'con un factor de autorreplicación logístico significativo, aplicados durante diez años. '
    'Estos valores no deben interpretarse como predicciones, sino como la consecuencia '
    'lógica del modelo bajo las hipótesis más favorables. La amplitud del rango entre '
    'escenarios (17.6× a 96.5×) refleja la incertidumbre sobre K y γ, los dos parámetros '
    'que determinan la magnitud y velocidad de la autorreplicación. Nótese que el factor ×2a '
    'del escenario Óptimo (1.8×) es menor que el del Base (2.0×): la curva logística '
    'concentra su aceleración alrededor de t₀ = 5 años, de modo que antes de ese punto '
    'el factor Φ_L aún no ha desplegado su potencial — propiedad característica de la S-curve.'
)

add_para(
    'La diferencia entre K=2 con γ=0.5 y K=6 con γ=1.5 no es de grado, es de naturaleza: '
    'el primero describe un ecosistema de IA que alcanza un multiplicador moderado con '
    'adopción gradual; el segundo describe un paradigma que acelera rápidamente hasta un '
    'multiplicador seis veces mayor. En ambos casos el crecimiento está acotado (Φ_L → 1+K '
    'cuando t → ∞), propiedad coherente con la teoría de difusión tecnológica de Bass (1969) '
    'y con los límites observados históricamente en la adopción de tecnologías de propósito general.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. FACTORES DE RIESGO
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Factores de Riesgo y Limitaciones del Modelo', level=1)

doc.add_heading('7.1 Riesgos de la tesis', level=2)

add_bullet(
    'Obsolescencia del cobre: si se valida industrialmente el uso de grafeno u otros '
    'nanomateriales como sustitutos del cobre a escala (TRL > 7), las posiciones en '
    'metales industriales perderían su fundamento. El umbral de alerta es la validación '
    'comercial, no el laboratorio.'
)
add_bullet(
    'Saturación energética: el anuncio de fusión nuclear neta comercialmente viable antes '
    'de 2030 cambiaría radicalmente el caso de inversión en uranio y utilities, '
    'al hacer obsoletos los reactores de fisión como fuente de energía basal.'
)
add_bullet(
    'Riesgo regulatorio sobre Bitcoin: legislación que limite la identidad de máquina '
    'en blockchain o que restrinja el uso de Bitcoin por agentes de IA afectaría '
    'directamente al pilar de escasez digital.'
)
add_bullet(
    'Deflación masiva de capex de IA: si el coste de hardware de inferencia cae >80% '
    'en dos años, la demanda de chips especializados podría colapsar, afectando al '
    'pilar de autorreplicación.'
)
add_bullet(
    'Concentración de mercado: la tesis asume competencia en el desarrollo de LLMs. '
    'Un escenario de monopolio regulado por un único actor estatal podría cambiar '
    'radicalmente los incentivos del ecosistema.'
)

doc.add_heading('7.2 Limitaciones del modelo matemático', level=2)

add_para(
    'El modelo V(t) = Capital × (1+r)^t × Φ_L(t) presenta varias limitaciones '
    'metodológicas que deben tenerse en cuenta al interpretar sus proyecciones:'
)

add_bullet(
    'Parámetro K parcialmente no calibrable desde datos de mercado: a diferencia de r '
    '(calibrable desde SOX) y γ (calibrable desde cadencia LLM), K representa el '
    'multiplicador máximo de autorreplicación en el estado maduro del ecosistema, lo que '
    'requiere hipótesis sobre el estado final de la IA que no son completamente '
    'objetivables. K se calibra actualmente desde ratios de expansión del TAM de IA, '
    'que son en sí mismos estimaciones con alta incertidumbre.'
)
add_bullet(
    'Sensibilidad al año de inflexión t₀: el valor a 10 años es especialmente sensible '
    'a la posición de t₀ relativa al horizonte de proyección. Si t₀ = 8 en lugar de t₀ = 5, '
    'la S-curve apenas ha completado su transición en el horizonte analizado, reduciendo '
    'los factores de multiplicación de forma sustancial. Esta sensibilidad debe comunicarse '
    'explícitamente en cualquier presentación de los escenarios.'
)
add_bullet(
    'Sin correlación entre activos: el modelo trata cada pilar de forma independiente '
    'y no modela la correlación entre categorías, lo que subestima el riesgo de '
    'caídas simultáneas en escenarios de stress sistémico.'
)
add_bullet(
    'Backtesting limitado: dado que varios de los activos del modelo (ETFs de quantum '
    'computing, smart grid infrastructure) tienen historiales de precios inferiores a '
    'cinco años, no es posible realizar un backtesting estadísticamente robusto.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 8. CONCLUSIONES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('8. Conclusiones', level=1)

add_para(
    'Este paper ha presentado el modelo Escasez y Resiliencia como un marco matemático '
    'tractable para invertir en la intersección de la escasez física y digital durante '
    'el período de despliegue masivo de la inteligencia artificial. Los principales '
    'hallazgos son:'
)

add_bullet(
    'El modelo V(t) = Capital × (1+r)^t × Φ_L(t) captura de forma parsimoniosa '
    'la doble naturaleza del crecimiento del ecosistema de IA: un componente de '
    'crecimiento compuesto calibrable desde datos de mercado (r_IA ≈ SOX trailing 12M / 100) '
    'y un factor logístico de autorreplicación Φ_L(t) = 1 + K / (1 + e^(−γ(t − t₀))) '
    'con parámetros interpretables: K (multiplicador máximo), γ (velocidad de adopción, '
    'calibrable desde la cadencia de releases LLM) y t₀ (año de inflexión, detectable '
    'via la señal SOX/Cobre).'
)
add_bullet(
    'La Correlación SOX/Cobre es el indicador de paradigma más relevante para la '
    'gestión táctica de la cartera: la Divergencia (SOX sube, Cobre baja) señala el '
    'momento en que el valor se transfiere de los activos físicos a los digitales.'
)
add_bullet(
    'Los cinco pilares de la cartera cubren de forma complementaria tanto los activos '
    'que la IA necesita hoy (cobre, uranio, chips) como los que se aprecian a medida '
    'que la IA se vuelve más autónoma (Bitcoin, robótica) y los que protegen ante '
    'escenarios adversos (renta variable global).'
)
add_bullet(
    'La revisión trimestral de los parámetros, documentando los criterios utilizados '
    'en cada ajuste, es esencial para mantener la coherencia y reproducibilidad '
    'del proceso de inversión a lo largo del tiempo.'
)

add_para(
    'La pregunta central que articula toda la tesis permanece abierta: ¿cuándo la IA '
    'dejará de necesitar hierro? El modelo propuesto no responde a esta pregunta, '
    'pero sí proporciona las herramientas para detectar cuándo la respuesta está '
    'comenzando a materializarse en los precios de mercado. Esa capacidad de detección '
    'temprana — más que la precisión de las proyecciones — es el valor principal '
    'del framework presentado.'
)

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCIAS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('Referencias', level=1)

refs = [
    # Modelos de difusión tecnológica — fundamento de Φ_L(t)
    'Bass, F. M. (1969). A new product growth for model consumer durables. '
    'Management Science, 15(5), 215–227. https://doi.org/10.1287/mnsc.15.5.215',

    'Rogers, E. M. (2003). Diffusion of Innovations (5th ed.). Free Press.',

    # Selección de cartera y valoración de activos — marco G11/G12
    'Markowitz, H. (1952). Portfolio selection. '
    'The Journal of Finance, 7(1), 77–91. https://doi.org/10.2307/2975974',

    'Fama, E. F. (1970). Efficient capital markets: A review of theory and empirical work. '
    'The Journal of Finance, 25(2), 383–417. https://doi.org/10.2307/2325486',

    # Cambio tecnológico y crecimiento endógeno — marco O33
    'Romer, P. M. (1990). Endogenous technological change. '
    'Journal of Political Economy, 98(5, Part 2), S71–S102. https://doi.org/10.1086/261725',

    'Brynjolfsson, E., & McAfee, A. (2014). The Second Machine Age: Work, Progress, and '
    'Prosperity in a Time of Brilliant Technologies. W. W. Norton & Company.',

    # Leyes de escala en modelos de lenguaje — calibración de γ
    'Kaplan, J., McCandlish, S., Henighan, T., Brown, T. B., Chess, B., Child, R., … Amodei, D. (2020). '
    'Scaling laws for neural language models. arXiv:2001.08361. '
    'https://arxiv.org/abs/2001.08361',

    # Bitcoin como activo de reserva de valor — pilar Escasez Digital
    'Nakamoto, S. (2008). Bitcoin: A peer-to-peer electronic cash system. '
    'Retrieved from https://bitcoin.org/bitcoin.pdf',

    'Yermack, D. (2015). Is Bitcoin a real currency? An economic appraisal. '
    'In D. L. K. Chuen (Ed.), Handbook of Digital Currency (pp. 31–43). Academic Press. '
    'https://doi.org/10.1016/B978-0-12-802117-0.00002-3',

    # Minerales críticos para electrificación e IA — pilar Escasez Física
    'International Energy Agency. (2021). The Role of Critical Minerals in Clean Energy Transitions. '
    'IEA. https://www.iea.org/reports/the-role-of-critical-minerals-in-clean-energy-transitions',

    # Innovación disruptiva — marco de referencia para t₀
    'Christensen, C. M. (1997). The Innovator\'s Dilemma: When New Technologies Cause Great '
    'Firms to Fail. Harvard Business School Press.',

    # Aprendizaje automático — referencia técnica general
    'Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press. '
    'Retrieved from https://www.deeplearningbook.org',

    # Cambio de paradigma y destrucción creativa
    'Schumpeter, J. A. (1942). Capitalism, Socialism and Democracy. Harper & Brothers.',
]

for ref in refs:
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent   = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    p.paragraph_format.space_after   = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(ref)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

# ══════════════════════════════════════════════════════════════════════════════
# NOTA DEL AUTOR
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Nota del Autor', level=1)
add_para(
    'Jose Antonio Vilar es profesional de Relaciones con Inversores (IR) con años de experiencia '
    'en el sector y estudiante del Grado en Ciencia de Datos de la Universitat Oberta de Catalunya '
    '(UOC). Este paper es un working paper de investigación personal que combina su experiencia '
    'práctica en mercados financieros con el rigor analítico de su formación académica en curso. '
    'Las proyecciones presentadas son el resultado de la aplicación del modelo matemático descrito '
    'y no constituyen asesoramiento financiero ni recomendación de inversión. El autor mantiene '
    'posiciones en los activos mencionados. Los resultados pasados no garantizan resultados futuros. '
    'Toda inversión conlleva riesgo de pérdida parcial o total del capital invertido.',
    italic=True, size=10
)

add_para(
    'El autor agradece las conversaciones con la comunidad inversora que han contribuido '
    'a refinar los argumentos presentados en este paper. Los errores que pudieran '
    'subsistir son responsabilidad exclusiva del autor.',
    italic=True, size=10
)

add_para(
    'Contacto: javsprivate@gmail.com',
    italic=True, size=10
)

# ── Guardar ───────────────────────────────────────────────────────────────────
out_path = '/sessions/exciting-clever-brown/mnt/Cartera-Inversion/exports/Vilar_2026_Escasez_Resiliencia_SSRN.docx'
doc.save(out_path)
print(f'Guardado: {out_path}')
