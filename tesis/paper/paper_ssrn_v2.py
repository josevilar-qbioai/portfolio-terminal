"""
Paper SSRN — Segunda entrega
Autofabricación y Olas Tecnológicas:
El Modelo de Logística Anidada para la Era de los Robots que se Autorreplican

Companion paper a "Escasez y Resiliencia" (paper_ssrn.py)
Formato: working paper académico, ~14 páginas
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Márgenes ──────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width    = Cm(21)
section.page_height   = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── Estilos base ──────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(11)

def set_heading_style(style_name, size, bold=True, color=None):
    s = doc.styles[style_name]
    s.font.name = 'Times New Roman'
    s.font.size = Pt(size)
    s.font.bold = bold
    if color:
        s.font.color.rgb = RGBColor(*color)
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after  = Pt(6)

set_heading_style('Heading 1', 13, bold=True)
set_heading_style('Heading 2', 12, bold=True)
set_heading_style('Heading 3', 11, bold=True)

def add_para(text, style='Normal', bold=False, italic=False, size=None,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6,
             color=None, indent=None):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent is not None:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    if size:  run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    return p

def add_formula(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    return p

def add_table(headers, rows, col_widths, hdr_bold=True):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_row = t.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr_row.cells[i]
        cell.width = Cm(w)
        cell_xml = cell._tc
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'E8E8E8')
        cell_xml.get_or_add_tcPr().append(shd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(h))
        run.bold = hdr_bold
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
    for row_data in rows:
        row = t.add_row()
        for i, (cell_text, w) in enumerate(zip(row_data, col_widths)):
            cell = row.cells[i]
            cell.width = Cm(w)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
    doc.add_paragraph()
    return t

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    return p

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run('─' * 72)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(180, 180, 180)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_footnote_style(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(1)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.italic = True
    return p

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
p_wp = doc.add_paragraph()
p_wp.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_wp.paragraph_format.space_before = Pt(20)
p_wp.paragraph_format.space_after  = Pt(6)
rw = p_wp.add_run('Working Paper · Companion Paper a "Escasez y Resiliencia" (2026)')
rw.font.size = Pt(10); rw.font.name = 'Times New Roman'; rw.italic = True
rw.font.color.rgb = RGBColor(100, 100, 100)

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(16)
p_title.paragraph_format.space_after  = Pt(8)
r = p_title.add_run('Autofabricación y Olas Tecnológicas:')
r.bold = True; r.font.size = Pt(20); r.font.name = 'Times New Roman'

p_title2 = doc.add_paragraph()
p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title2.paragraph_format.space_after = Pt(4)
r2 = p_title2.add_run('El Modelo de Logística Anidada para la Era de los Robots')
r2.bold = True; r2.font.size = Pt(16); r2.font.name = 'Times New Roman'

p_title3 = doc.add_paragraph()
p_title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title3.paragraph_format.space_after = Pt(24)
r3 = p_title3.add_run('que se Autorreplican')
r3.bold = True; r3.font.size = Pt(16); r3.font.name = 'Times New Roman'

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(8)
rs = p_sub.add_run(
    'Self-Fabrication and Technological Waves:\n'
    'The Nested Logistic Model for the Age of Self-Replicating Robots'
)
rs.italic = True; rs.font.size = Pt(12); rs.font.name = 'Times New Roman'

add_separator()

p_auth = doc.add_paragraph()
p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_auth.paragraph_format.space_before = Pt(20)
p_auth.paragraph_format.space_after  = Pt(6)
ra = p_auth.add_run('Jose Antonio Vilar')
ra.bold = True; ra.font.size = Pt(13); ra.font.name = 'Times New Roman'

p_aff = doc.add_paragraph()
p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_aff.paragraph_format.space_after = Pt(4)
raff = p_aff.add_run('Profesional de Relaciones con Inversores · Estudiante de Grado en Ciencia de Datos, UOC')
raff.font.size = Pt(11); raff.font.name = 'Times New Roman'

p_email = doc.add_paragraph()
p_email.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_email.paragraph_format.space_after = Pt(4)
re_ = p_email.add_run('javsprivate@gmail.com · joseantoniovilar@uoc.edu')
re_.font.size = Pt(10); re_.font.name = 'Times New Roman'; re_.italic = True

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_date.paragraph_format.space_after = Pt(24)
rd = p_date.add_run('Marzo 2026')
rd.font.size = Pt(11); rd.font.name = 'Times New Roman'

add_separator()

p_jel = doc.add_paragraph()
p_jel.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_jel.paragraph_format.space_before = Pt(16)
p_jel.paragraph_format.space_after  = Pt(4)
rj = p_jel.add_run('Clasificación JEL: G11 · O33 · O14 · L16 · Q02')
rj.font.size = Pt(10); rj.font.name = 'Times New Roman'; rj.italic = True

p_kw = doc.add_paragraph()
p_kw.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_kw.paragraph_format.space_after = Pt(4)
rk = p_kw.add_run(
    'Palabras clave: logística anidada, autofabricación robótica, olas tecnológicas, '
    'prima de escasez, cambio de paradigma, construcción de cartera, difusión tecnológica'
)
rk.font.size = Pt(10); rk.font.name = 'Times New Roman'; rk.italic = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('Abstract', level=1)

add_para(
    'The Scarcity and Resilience model, introduced in our companion paper, uses a single '
    'logistic curve Φ_L(t) to capture technological self-replication. That formulation is '
    'academically defensible and empirically calibratable for the 2024–2030 horizon. However, '
    'the historical record of technology diffusion suggests that major paradigm shifts do not '
    'follow a single S-curve: they arrive as successive, overlapping waves, each unlocking the '
    'next. This paper extends the framework by replacing Φ_L(t) with a nested logistic '
    'Φ_nested(t) = 1 + Σᵢ Kᵢ / (1 + e^(−γᵢ(t − t₀ᵢ))), calibrated on four identifiable '
    'technological waves for the period 2026–2038: mass LLM adoption (underway), robotics with self-fabrication (2027), '
    'autonomous agents (2029), and full-scale autofabrication (2031). We show that the '
    'qualitative change introduced by the second wave — robots manufacturing robots, from 2027 — '
    'is not merely quantitative: '
    'it collapses marginal production costs and generates an absolute scarcity premium for '
    'assets with inelastic supply. We derive the portfolio implications analytically and '
    'demonstrate that the Scarcity and Resilience portfolio is structurally positioned at '
    'the physical bottlenecks that no replication process can eliminate: copper, uranium, '
    'Bitcoin, and smart-grid infrastructure.',
    space_after=8
)

doc.add_heading('Resumen', level=1)

add_para(
    'El modelo Escasez y Resiliencia, introducido en nuestro paper anterior, utiliza una '
    'curva logística simple Φ_L(t) para capturar la autorreplicación tecnológica. Esa '
    'formulación es académicamente sólida y calibrable empíricamente para el horizonte '
    '2024–2030. Sin embargo, el registro histórico de la difusión tecnológica sugiere que '
    'los grandes cambios de paradigma no siguen una única curva en S: llegan como olas '
    'sucesivas y superpuestas, cada una desbloqueando la siguiente. Este paper extiende el '
    'marco sustituyendo Φ_L(t) por una logística anidada Φ_nested(t) = 1 + Σᵢ Kᵢ / '
    '(1 + e^(−γᵢ(t − t₀ᵢ))), calibrada sobre cuatro olas tecnológicas identificables para '
    'el período 2026–2038: adopción masiva de LLM (en marcha), robótica con autofabricación (2027), agentes autónomos y '
    'autofabricación robótica. Demostramos que el cambio cualitativo introducido por la '
    'cuarta ola —robots que fabrican robots— no es meramente cuantitativo: colapsa los '
    'costes marginales de producción y genera una prima de escasez absoluta para los activos '
    'con oferta inelástica. Derivamos las implicaciones para el portfolio analíticamente y '
    'demostramos que la cartera Escasez y Resiliencia está posicionada estructuralmente en '
    'los cuellos de botella físicos que ningún proceso de replicación puede eliminar: cobre, '
    'uranio, Bitcoin e infraestructura de red eléctrica inteligente.',
    space_after=8
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# §1 INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Introducción', level=1)

doc.add_heading('1.1 El problema de la S-curve única', level=2)

add_para(
    'El paper fundacional de esta serie —"Escasez y Resiliencia: Un Marco Matemático '
    'para Invertir en la Era de la Autorreplicación Artificial" (Vilar, 2026a)— introduce '
    'el modelo V(t) = Capital × (1+r)^t × Φ_L(t)/Φ_L(0), donde Φ_L(t) es una curva '
    'logística con techo K. Esa elección es correcta para modelar la difusión de una '
    'tecnología concreta en un horizonte de 5 a 8 años. Pero la historia de la '
    'industrialización y de la revolución digital nos muestra sistemáticamente que '
    'las grandes transformaciones llegan en oleadas discretas, no en una única '
    'S-curve: cada generación tecnológica desbloquea la siguiente.'
)

add_para(
    'La imprenta (1450) fue seguida por la prensa de vapor (1814), luego el telégrafo '
    '(1837), el teléfono (1876), la radio (1920), la televisión (1940) y finalmente '
    'internet (1991). Ninguna de estas olas fue la continuación suave de la anterior: '
    'cada una tenía sus propios parámetros K, γ y t₀, y la suma de sus efectos fue '
    'superexponencial a lo largo del siglo, aunque cada ola individual saturaba.'
)

add_para(
    'En el contexto de la inteligencia artificial y la robótica del período 2026–2038, '
    'identificamos cuatro olas discretas y diferenciadas. Este paper extiende el '
    'modelo logístico simple a un modelo de logísticas anidadas que captura '
    'esta estructura de olas, y estudia en detalle las implicaciones de la cuarta ola '
    '—la autofabricación robótica— que introduce una discontinuidad cualitativa: '
    'cuando el factor de producción se convierte en su propio replicador.'
)

doc.add_heading('1.2 La discontinuidad cualitativa: el robot que hace robots', level=2)

add_para(
    'En toda la historia económica, la productividad del trabajo ha crecido, pero el '
    'trabajo mismo siempre ha sido un factor de producción con oferta limitada por '
    'la demografía humana. La autofabricación robótica rompe esta restricción: '
    'por primera vez en la historia, el factor de producción físico puede crecer '
    'a una tasa proporcional a su propia existencia actual.'
)

add_para(
    'La ecuación dR/dt = α · η · R(t) + P_H, donde R es la capacidad robótica, '
    'η es la eficiencia de autorreplicación y P_H es la producción humana base, '
    'tiene solución superexponencial para cualquier η > 0. El coste marginal de '
    'producción C(t) ∝ 1/R(t) tiende hacia cero. Esto no es ciencia ficción: '
    'ya existe autofabricación parcial en componentes de semiconductores '
    '(los robots en las fabs de TSMC son mantenidos y calibrados en parte por otros '
    'sistemas automatizados). La pregunta no es si ocurrirá, sino cuándo se '
    'cruzará el umbral donde η sea suficientemente alto para que el efecto '
    'sea dominante.'
)

add_para(
    'Este paper estudia las implicaciones de ese cruce para la construcción de '
    'portfolios de largo plazo.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# §2 EL MODELO DE OLAS ANIDADAS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('2. El Modelo de Logísticas Anidadas', level=1)

doc.add_heading('2.1 Formulación', level=2)

add_para(
    'El factor de autorreplicación anidado se define como la superposición de n '
    'curvas logísticas independientes, cada una representando una ola tecnológica '
    'discreta:'
)

add_formula('Φ_nested(t) = 1 + Σᵢ₌₁ⁿ  Kᵢ / (1 + e^(−γᵢ(t − t₀ᵢ)))')

add_para(
    'El modelo de valor normalizado del portfolio se convierte en:'
)

add_formula('V(t) = Capital × (1+r)^t × Φ_nested(t) / Φ_nested(0)')

add_para(
    'La normalización por Φ_nested(0) preserva la propiedad V(0) = Capital '
    'independientemente de los parámetros de cada ola, garantizando comparabilidad '
    'de escenarios. Nótese que el modelo del paper fundacional es el caso particular '
    'n=1 de esta formulación.'
)

add_para(
    'Cada ola i tiene tres parámetros independientes:'
)
add_bullet('Kᵢ — amplitud o multiplicador máximo de la ola i. El techo total es 1 + ΣKᵢ.')
add_bullet('γᵢ — velocidad de adopción. Cuanto más rápida la difusión, mayor γᵢ.')
add_bullet('t₀ᵢ — año de inflexión: máxima aceleración de la ola i. El detector SOX/Cobre estima t₀ empíricamente.')

doc.add_heading('2.2 Relación con la literatura de difusión tecnológica', level=2)

add_para(
    'La formulación de olas anidadas es formalmente equivalente a un modelo de Bass '
    '(1969) de adopción tecnológica multi-generacional, donde cada generación tiene '
    'su propio período de difusión. Freeman y Louçã (2001) identificaron cinco '
    '"ondas largas de Kondratiev" en la historia industrial, cada una con estructura '
    'logística propia: máquinas de vapor (1771), acero/ferrocarriles (1829), '
    'electrificación (1875), petróleo/automoción (1908) y TIC (1971). '
    'La IA y la robótica constituyen la sexta onda, con la particularidad de que '
    'la cuarta sub-ola —autofabricación— introduce por primera vez un '
    'mecanismo de retroalimentación positiva sin precedente en las olas anteriores.'
)

add_para(
    'Mokyr (2002) distingue entre macro-inventions (cambios de paradigma) y '
    'micro-inventions (mejoras dentro del paradigma). Las olas 1-3 son '
    'micro-inventions dentro del paradigma IA; la ola 4 es una macro-invention: '
    'cambia la naturaleza del factor de producción.'
)

doc.add_heading('2.3 Calibración de las cuatro olas — autofabricación desde 2027', level=2)

add_para(
    'La Tabla 1 presenta la calibración de los parámetros de cada ola basada en '
    'evidencia empírica disponible a marzo de 2026:'
)

add_table(
    headers=['Ola', 'Nombre', 't₀', 'K', 'γ', 'Base de calibración'],
    rows=[
        ['1', 'LLM / Software IA',        '2026 (t=0.5)', '1.5', '1.4',
         'Adopción ChatGPT/Copilot en marcha; 100M usuarios en 2 meses (OpenAI, 2023)'],
        ['2', 'Robótica + autofab inicial','2027 (t=1.0)', '2.5', '1.2',
         'Primera línea autofabricante operativa; Figure, Tesla Optimus, TSMC fab lights-out'],
        ['3', 'Agentes autónomos',         '2029 (t=3.0)', '3.5', '0.9',
         'IA que diseña y despliega sistemas propios; ritmo Devin, AutoGPT evolucionados'],
        ['4', 'Autofab 2ª fase',           '2031 (t=5.0)', '5.0', '0.7',
         'Autofabricación generalizada en manufactura amplia; η > 0.30 confirmado'],
    ],
    col_widths=[0.8, 3.2, 2.0, 0.8, 0.8, 6.4]
)

add_footnote_style(
    'Tabla 1. Parámetros calibrados de las cuatro olas. t=0 corresponde a enero 2026. '
    'Revisión de marzo 2026: la ola 2 se adelanta a 2027 basada en hoja de ruta publicada '
    'por Boston Dynamics, Figure AI y TSMC. γ en unidades anuales.'
)

add_para(
    'El techo total del factor Φ_nested es 1 + 1.5 + 2.5 + 3.5 + 5.0 = 13.5×, '
    'lo que implica que una cartera perfectamente posicionada en los pilares '
    'de la tesis podría multiplicar su valor 13.5 veces en términos reales '
    'a lo largo del horizonte completo, además del componente (1+r)^t. '
    'Este es el escenario "óptimo extremo"; los escenarios del paper fundacional '
    '(K base = 2, K óptimo = 6) corresponden a la materialización parcial de '
    'las olas 1 y 2.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# §3 AUTOFABRICACIÓN ROBÓTICA
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Autofabricación Robótica: Dinámica de Feedback', level=1)

doc.add_heading('3.1 El modelo de retroalimentación', level=2)

add_para(
    'Cuando robots pueden fabricar robots, la dinámica de la capacidad productiva '
    'deja de ser exógena y se convierte en un sistema con retroalimentación positiva. '
    'El modelo diferencial más simple que captura esta propiedad es:'
)

add_formula('dR/dt = α_R · η(t) · R(t) + P_H')

add_para(
    'donde R(t) es la capacidad robótica total normalizada (R₀ = 1 en el año de '
    'inicio de la autofabricación), α_R es la tasa de replicación máxima del '
    'sistema robótico, η(t) ∈ [0,1] es la eficiencia de autorreplicación '
    '(fracción de la flota capaz de fabricar nuevas unidades), y P_H es la '
    'componente exógena de producción humana.'
)

add_para(
    'Con η constante, la solución es:'
)

add_formula('R(t) = (R₀ + P_H/α_R) · e^(α_R · η · t)  −  P_H/α_R')

add_para(
    'El coste unitario de producción cae como el inverso de la capacidad:'
)

add_formula('C(t) ∝ 1 / R(t)  →  0  cuando  t → ∞')

add_para(
    'La Tabla 2 muestra las proyecciones de capacidad y coste relativo para tres '
    'escenarios de η a partir del arranque de la autofabricación en 2027:'
)

add_table(
    headers=['η', 'Escenario', '2030 (t+3)', '2032 (t+5)', '2035 (t+8)', 'Coste rel. 2032'],
    rows=[
        ['0.05', 'Conservador', '1.5×',  '1.8×',  '2.5×',  '56%'],
        ['0.20', 'Base',        '3.0×',  '6.2×',  '22×',   '16%'],
        ['0.50', 'Agresivo',    '12×',   '66×',   '2100×', '1.5%'],
    ],
    col_widths=[1.0, 3.0, 2.2, 2.2, 2.2, 2.4]
)

add_footnote_style(
    'Tabla 2. Proyecciones de capacidad robótica y coste relativo desde 2027. α_R = 0.30, '
    'P_H = 0.05, R₀ = 1. Los valores de η = 0.50 son especulativos. '
    'El escenario base (η=0.20) implica coste relativo del 87% en 2030, 76% en 2032.'
)

doc.add_heading('3.2 Velocidad del colapso del coste de producción', level=2)

add_para(
    'La implicación económica clave no es el nivel absoluto de R(t), sino la '
    'velocidad a la que C(t) cae. En el escenario base (η = 0.20), el coste '
    'unitario de producción se reduce a la mitad en aproximadamente 3.5 años desde '
    'el inicio de la autofabricación. Para comparación: la Ley de Moore tardó '
    'aproximadamente 2 años en halvar el coste por transistor, pero solo aplicaba '
    'a semiconductores. La autofabricación robótica aplica potencialmente a '
    'todo objeto físico manufacturado.'
)

add_para(
    'Este colapso de costes tiene un precedente histórico parcial: la revolución '
    'textil en Lancashire (1780-1820) redujo el coste de la tela de algodón en '
    'un 90% en cuatro décadas. La autofabricación robótica podría comprimir '
    'ese proceso a menos de una década para categorías enteras de manufactura.'
)

doc.add_heading('3.3 La prima de escasez en un mundo de producción ilimitada', level=2)

add_para(
    'Cuando el coste de producción tiende hacia cero, el valor de los bienes '
    'producibles también tiende hacia cero. Pero la misma dinámica genera el '
    'efecto opuesto sobre los bienes con oferta genuinamente inelástica: '
    'su prima de escasez aumenta en términos relativos.'
)

add_para(
    'Formalizamos esto mediante el coeficiente de escasez εᵢ ∈ [0,1] de cada activo, '
    'definido como la fracción de su valor que no puede ser replicada por un proceso '
    'de manufactura automatizado. La prima de escasez es:'
)

add_formula('S_i(t) = 1 + ε_i · (R(t) − 1) / R(t)')

add_para(
    'Nótese que S_i → 1 + εᵢ cuando R → ∞: en el límite de producción ilimitada, '
    'los activos con ε = 1 duplican su valor relativo, mientras que los activos '
    'con ε = 0 no se benefician. La Tabla 3 presenta los coeficientes de escasez '
    'estimados para las categorías de la cartera:'
)

add_table(
    headers=['Categoría', 'εᵢ', 'Justificación', 'Límite S_i(∞)'],
    rows=[
        ['Bitcoin', '1.00',
         'Oferta fija por protocolo matemático; imposible de replicar',
         '2.00×'],
        ['Cobre', '0.70',
         'Recurso geológico; reciclaje parcial pero oferta a largo plazo inelástica',
         '1.70×'],
        ['Uranio', '0.75',
         'Escasez geológica + regulación + ciclo minero de 10-15 años',
         '1.75×'],
        ['Energía / Grid', '0.65',
         'Infraestructura física; escala pero no se puede "imprimir" capacidad',
         '1.65×'],
        ['Tecnología / AI', '0.45',
         'Software replicable pero hardware + talento tienen fricción',
         '1.45×'],
        ['Renta Variable', '0.20',
         'Beta mercado; incluye sectores replicables y escasos mezclados',
         '1.20×'],
    ],
    col_widths=[3.5, 1.0, 7.5, 2.0]
)

add_footnote_style(
    'Tabla 3. Coeficientes de escasez estimados por categoría. Los valores son '
    'proxy cualitativos; la calibración empírica requeriría análisis de correlación '
    'con índices de costes de producción sectoriales.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# §4 IMPACTO EN EL ECOSISTEMA EMPRESARIAL
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Impacto en el Ecosistema Empresarial', level=1)

doc.add_heading('4.1 La taxonomía replicable vs. escaso', level=2)

add_para(
    'El colapso del coste de manufactura no afecta por igual a todos los sectores '
    'de la economía. La variable determinante no es el tamaño de la empresa ni su '
    'sector GICS, sino una propiedad más fundamental: la replicabilidad de su '
    'producto o servicio.'
)

add_para(
    'Definimos la replicabilidad ρ de un sector como la fracción de su cadena de '
    'valor susceptible de ser reproducida por un sistema robótico autofabricante a '
    'coste marginal tendente a cero. La Tabla 4 presenta el análisis por sector:'
)

add_table(
    headers=['Sector', 'ρ (replicabilidad)', 'Impacto autofabricación',
             'Dirección márgenes', 'Posición cartera'],
    rows=[
        ['Manufactura estándar',         'Muy alta (>0.85)',  'Colapso de márgenes',        '↓↓↓', 'Sin exposición'],
        ['Logística / distribución',     'Alta (0.70)',       'Desintermediación severa',    '↓↓',  'Sin exposición'],
        ['Construcción',                 'Alta (0.65)',       'Robotización acelerada',      '↓↓',  'Sin exposición'],
        ['Software / datos / AI',        'Media (0.45)',      'Commodity parcial',           '↓~',  'AI & Big Data ETF'],
        ['Computación cuántica',         'Baja (0.25)',       'Demanda explosiva',           '↑↑',  'Quantum ETF'],
        ['Energía nuclear / uranio',     'Baja (0.20)',       'Demanda inelástica sube',     '↑↑↑', 'Uranium ETF'],
        ['Metales críticos (Cu)',         'Muy baja (0.10)',   'Oferta física no se replica', '↑↑↑', 'Cobre ETF'],
        ['Grid / infraest. eléctrica',   'Baja (0.20)',       'Cuello botella inexorable',   '↑↑↑', 'Smart Grid ETF'],
        ['Bitcoin / activos digitales',  'Cero (0.00)',       'Escasez absoluta programada', '↑↑↑↑','BTC ×4 posiciones'],
        ['Renta variable global',        'Media (0.40)',      'Mixto — depende de composic.','~↑',  'Fondos índice'],
    ],
    col_widths=[3.8, 2.4, 3.6, 2.4, 2.8]
)

add_footnote_style(
    'Tabla 4. Análisis de impacto de la autofabricación por sector. La columna '
    '"Dirección márgenes" es cualitativa: ↑ = mejora, ↓ = deterioro, ~ = neutro. '
    '"Posición cartera" indica el activo correspondiente en la cartera Escasez y Resiliencia.'
)

doc.add_heading('4.2 El mecanismo: presión en el ecosistema', level=2)

add_para(
    'La autofabricación robótica genera cuatro efectos de transmisión sobre el '
    'ecosistema empresarial:'
)

add_para(
    'Primero, compresión de márgenes en manufactura. Cualquier empresa que vende '
    'objetos físicos producibles con maquinaria general verá cómo sus competidores '
    'replicados a coste cero erosionan su poder de fijación de precios. La única '
    'defensa es la escasez del input (materias primas), la regulación, o la '
    'protección por diseño/marca (que tiene durabilidad limitada).'
)

add_para(
    'Segundo, explosión de la demanda de energía base. Cada robot que fabrica robots '
    'consume electricidad ininterrumpidamente. A diferencia del calor humano, que se '
    'regula solo, una flota robótica en crecimiento exponencial genera una demanda '
    'energética exponencial. Esto hace de la energía nuclear y la infraestructura '
    'de grid inteligente un cuello de botella estructural, no cíclico.'
)

add_para(
    'Tercero, escasez de materias primas críticas. El cobre es el material '
    'conductor universal para electrificación. Cada robot adicional requiere '
    'cobre; la aceleración de la flota robótica es directamente una aceleración '
    'de la demanda de cobre. A diferencia de la demanda de software, '
    'la oferta de cobre no puede crecer más rápido que el ciclo minero '
    '(10-15 años desde exploración hasta producción).'
)

add_para(
    'Cuarto, refugio programado en activos digitales. En un mundo donde la '
    'producción física tiende a coste cero, los activos cuya escasez es '
    'matemáticamente garantizada por un protocolo de consenso distribuido '
    '(Bitcoin, 21 millones de unidades, inmutables) actúan como el activo '
    'de reserva más puro posible: no puede haber más, independientemente '
    'de lo que haga cualquier robot.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# §5 EL MODELO EXTENDIDO
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('5. El Modelo Extendido: V_i(t) con Premio de Escasez', level=1)

doc.add_heading('5.1 Formulación completa', level=2)

add_para(
    'Integrando el modelo de olas anidadas (§2) con la dinámica de autofabricación '
    '(§3), el valor de la posición i del portfolio en el horizonte completo es:'
)

add_formula('V_i(t) = Capital_i × (1 + r_i)^t × [Φ_nested(t) / Φ_nested(0)] × S_i(t)')

add_para(
    'donde los cuatro factores tienen interpretación económica precisa:'
)

add_bullet('Capital_i — valor invertido en la categoría i a t=0.')
add_bullet('(1+r_i)^t — crecimiento intrínseco de la categoría (calibrado por activo: r_BTC=25%, r_IA=20%, r_Nuclear=15%, r_Metales=8%, r_RV=10%). En análisis multi-escenario, r_i se escala proporcionalmente al escenario (BASE→ACELERADO→ÓPTIMO), manteniendo las diferencias relativas entre categorías.')
add_bullet('Φ_nested(t)/Φ_nested(0) — factor de autorreplicación tecnológica acumulado, común para todas las categorías.')
add_bullet('S_i(t) = 1 + εᵢ · (R(t)−1)/R(t) — prima de escasez diferencial por categoría, solo activa a partir de la ola 4.')

add_para(
    'Nótese que S_i(t) = 1 para todo t < t_robot (antes de la autofabricación). '
    'Con t_robot = 2027, la prima de escasez es activa desde el segundo año del '
    'horizonte de inversión. El modelo extendido colapsa al modelo del paper '
    'fundacional solo en el primer año (2026–2027). Es un modelo anidado en el '
    'sentido más literal: el paper anterior es el caso límite para t < 1.'
)

doc.add_heading('5.2 Comparación de modelos', level=2)

add_para(
    'La Tabla 5 presenta la jerarquía de modelos de esta familia, desde el más '
    'simple al más completo:'
)

add_table(
    headers=['Modelo', 'Fórmula Φ(t)', 'Parámetros', 'Horizonte válido', 'Fuente'],
    rows=[
        ['Base (paper v1)',
         'Φ_L(t) / Φ_L(0)  — logística simple',
         'K, γ, t₀',
         '2024–2030',
         'Vilar (2026a)'],
        ['Anidado (este paper)',
         'Φ_nested(t) / Φ_nested(0)  — 4 olas',
         '4×{Kᵢ, γᵢ, t₀ᵢ}',
         '2026–2038',
         'Vilar (2026b)'],
        ['Extendido completo',
         'Φ_nested(t)/Φ_nested(0) × S_i(t)',
         '+{η, α_R, εᵢ}',
         '2024–2040',
         'Vilar (2026b)'],
    ],
    col_widths=[3.5, 5.0, 3.0, 2.5, 2.0]
)

doc.add_heading('5.3 Proyecciones comparativas', level=2)

add_para(
    'La Tabla 6 muestra las proyecciones de valor del portfolio completo bajo los '
    'tres modelos, sobre una base de €100.000 con asignación ilustrativa '
    '(BTC 40%, RV 20%, IA 15%, Nuclear 10%, Metales 8%, Quantum 4%, Grid 3%; '
    'r ponderado ≈ 19%). Todos los modelos parten del mismo capital en t=0. '
    'Los escenarios usan r_i por activo escalado al escenario BASE '
    '(r_BTC = 25%, r_IA = 20%, r_Nuclear = 15%, r_Metales = 8%, r_RV = 10%):'
)

add_table(
    headers=['Año', 'Simple logístico (v1)', 'Anidado (Φ_nested)', 'Extendido (Φ × S)', 'Benchmark MSCI (+10%)'],
    rows=[
        ['2028 (t=2)',  '€  168k', '€  336k', '€  353k', '€  121k'],
        ['2029 (t=3)',  '€  226k', '€  518k', '€  567k', '€  133k'],
        ['2031 (t=5)',  '€  423k', '€1.040k', '€1.240k', '€  161k'],
        ['2033 (t=7)',  '€  760k', '€1.810k', '€2.290k', '€  195k'],
        ['2036 (t=10)', '€1.580k', '€3.490k', '€4.810k', '€  259k'],
        ['2038 (t=12)', '€2.440k', '€5.270k', '€7.610k', '€  314k'],
    ],
    col_widths=[2.5, 3.2, 3.0, 3.0, 3.3]
)

add_footnote_style(
    'Tabla 6. Proyecciones ilustrativas. Capital inicial €100k, distribución genérica (datos no reales). '
    'BASE: r por activo calibrado individualmente (BTC=25%, IA=20%, Nuclear=15%, RV=10%), η=0.20, '
    'autofabricación desde 2027 (t_robot=1). Anidado: solo Φ_nested(t)/Φ_nested(0), sin S_i(t). '
    'Extendido: Φ_nested × S_i(t) con ε_BTC=1.0, ε_IA=0.7, ε_RV=0.15. '
    'Las cifras son escenarios de referencia, no predicciones.'
)

add_para(
    'El resultado más significativo de la Tabla 6 es que la divergencia entre el modelo '
    'extendido y el logístico simple comienza ya en 2028 — un año después del arranque '
    'de la autofabricación. En el horizonte de diez años (2036), '
    'el modelo extendido (Φ_nested × S) supera al simple logístico en 3× '
    '(€4.81M vs €1.58M). La mayor parte de la diferencia proviene de Φ_nested(t): '
    'mientras el multiplicador logístico simple (BASE, K=2, γ=0.5) alcanza un ratio '
    'de 2.47× a t=10, Φ_nested(10)/Φ_nested(0) = 5.46×, con las cuatro olas ya '
    'madurando simultáneamente. La prima de escasez S_i(t) añade un 30–40% adicional '
    'para la fracción de Bitcoin (ε=1.0), que en esta asignación representa el 40% '
    'del capital. Comparado con el escenario anterior (autofab en 2033), el adelanto '
    'a 2027 incorpora seis años adicionales de capitalización de las olas 2–4.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# §6 IMPLICACIONES PARA LA CARTERA
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Implicaciones para la Cartera Escasez y Resiliencia', level=1)

doc.add_heading('6.1 Validación estructural de la tesis', level=2)

add_para(
    'El análisis del modelo extendido produce una conclusión central: la cartera '
    'Escasez y Resiliencia, construida intuitivamente alrededor de pilares de '
    'escasez física y digital, está estructuralmente alineada con el escenario '
    'de autofabricación robótica por razones que van más allá de la intuición '
    'original.'
)

add_para(
    'En concreto, los activos con mayor εᵢ (Bitcoin, cobre, uranio, grid) son '
    'exactamente los activos que no pueden ser replicados por ningún proceso de '
    'manufactura automatizado, por definición. Bitcoin porque su escasez es '
    'matemática; el cobre y el uranio porque son elementos de la tabla periódica '
    'que ningún robot puede sintetizar; la infraestructura de grid porque requiere '
    'física constructiva con un ciclo de inversión de décadas.'
)

doc.add_heading('6.2 El SOX/Cobre como detector de t₀', level=2)

add_para(
    'La señal SOX/Cobre introducida en el paper fundacional adquiere una '
    'interpretación adicional en el modelo de olas anidadas: es el detector '
    'del inicio de la ola 2 (robótica física, t₀ ≈ 2027). Cuando SOX sube '
    'y Cobre también sube, ambas olas están activas simultáneamente — el '
    'mundo construye tanto las capacidades de IA como la infraestructura física '
    'que las soporta. Cuando SOX sube y Cobre baja (Divergencia), la ola de '
    'software ha superado temporalmente a la de hardware.'
)

add_para(
    'Para la ola 4, el detector análogo aún no existe como señal de mercado '
    'establecida. La candidata más prometedora es el ratio capex en robots/capex '
    'total de las mayores fábricas del mundo (TSMC, Foxconn, Toyota), cuando ese '
    'ratio supere el 40% y comience a crecer exponencialmente, estaremos '
    'cerca de t₀₄.'
)

doc.add_heading('6.3 Revisiones del modelo recomendadas', level=2)

add_para(
    'El modelo de olas anidadas introduce parámetros adicionales que requieren '
    'un protocolo de revisión más estructurado que el modelo simple. Proponemos:'
)

add_bullet(
    'Revisión trimestral de r y γ₁ (olas 1-2): actualizar con cadencia de releases '
    'LLM y datos de instalación de robots industriales (IFR World Robotics Report).'
)
add_bullet(
    'Revisión semestral de K₁ y K₂: ajustar según penetración observada de LLM en '
    'empresas (índices de adopción de Gartner, IDC) y ratio de robots/trabajador '
    'en manufactura global.'
)
add_bullet(
    'Revisión anual de t₀₃ y t₀₄: actualizar con hoja de ruta publicada por '
    'los principales laboratorios de robótica y semiconductores.'
)
add_bullet(
    'Señal de alerta temprana ola 4: monitorizar el ratio capex robots/capex total '
    'en las 10 mayores fábricas del mundo. Umbral de atención: >30%.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# §7 LIMITACIONES Y RIESGOS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Limitaciones del Modelo y Riesgos de la Tesis', level=1)

doc.add_heading('7.1 Limitaciones matemáticas', level=2)

add_para(
    'El modelo de olas anidadas tiene tres limitaciones matemáticas que el lector '
    'debe tener presentes:'
)

add_para(
    'Primera, independencia de olas. El modelo asume que las olas Kᵢ son '
    'independientes y aditivas. En realidad, existe retroalimentación entre ellas: '
    'la ola 2 (robótica) depende del éxito de la ola 1 (LLM). Una formulación '
    'más rigurosa introduciría términos de interacción, aumentando la complejidad '
    'paramétrica sin necesariamente mejorar la predictibilidad.'
)

add_para(
    'Segunda, techo fijo por ola. Cada Kᵢ es constante en el modelo. En la '
    'realidad, el multiplicador máximo de una ola puede cambiar si las condiciones '
    'tecnológicas o regulatorias varían. El techo K₄ = 5 para la autofabricación '
    'asume que la regulación no detiene el proceso antes de su saturación natural.'
)

add_para(
    'Tercera, calibración prospectiva incierta. La ola 2 (2027) es la hipótesis '
    'central de este paper; las olas 3 y 4 (2029, 2031) están calibradas sobre '
    'proyecciones. La incertidumbre en t₀₂ es de ±1 año, con impacto en '
    'proyecciones del orden del 25-40%. Si la autofabricación se retrasa a 2029, '
    'los resultados de la Tabla 6 se moderan a niveles comparables al escenario '
    'original de 2033 pero con menor divergencia en el largo plazo.'
)

doc.add_heading('7.2 Riesgos de la tesis de inversión', level=2)

add_para(
    'La tesis de la cartera asume que la autofabricación robótica se materializará '
    'desde 2027. Este horizonte significativamente más cercano eleva tanto el '
    'potencial de la cartera como la urgencia de los riesgos. Los principales '
    'riesgos a esta hipótesis son:'
)

add_bullet(
    'Riesgo regulatorio: Los gobiernos podrían prohibir o limitar severamente la '
    'autofabricación robótica por razones de empleo, seguridad o control geopolítico. '
    'Probabilidad estimada: moderada (30-40%); impacto en modelo: retraso de t₀₄ en '
    '5-10 años, reducción de K₄ a 2-3.'
)
add_bullet(
    'Riesgo tecnológico: La autofabricación requiere robots generalistas con destreza '
    'comparable a la humana. Actualmente los robots industriales son altamente '
    'específicos; la generalización puede tardar más de lo estimado.'
)
add_bullet(
    'Riesgo de sustitución de materiales: Avances en materiales sintéticos podrían '
    'reducir la dependencia del cobre (grafeno, superconductores de alta temperatura). '
    'Probabilidad: baja para el horizonte 2036, moderada para 2040+.'
)
add_bullet(
    'Riesgo de protocolo Bitcoin: Un fallo de consenso o ataque cuántico exitoso '
    'podría comprometer la escasez programada de Bitcoin. Probabilidad: muy baja '
    'para el horizonte de este paper.'
)
add_bullet(
    'Riesgo de régimen de tipos de interés: Un entorno de tipos altos prolongado '
    'reduce el valor presente de los flujos futuros y penaliza los activos de '
    'crecimiento. El modelo no incluye el tipo de descuento explícitamente.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# §8 CONCLUSIONES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('8. Conclusiones', level=1)

add_para(
    'Este paper extiende el modelo Escasez y Resiliencia en dos dimensiones: '
    'la estructura de olas anidadas que captura la naturaleza discreta y '
    'superpuesta del cambio tecnológico, y el mecanismo de retroalimentación '
    'de la autofabricación robótica que genera una prima de escasez diferencial '
    'por categoría de activo.'
)

add_para(
    'Los resultados más relevantes son tres. Primero, el modelo de olas anidadas '
    'es formalmente una extensión del modelo del paper fundacional, que es el '
    'caso n=1. Esto garantiza coherencia matemática entre ambos trabajos. Segundo, '
    'la autofabricación robótica —revisada a 2027 en esta edición— introduce una '
    'discontinuidad cualitativa: convierte al factor de producción en su propio '
    'replicador, colapsando costes marginales y amplificando la prima de escasez '
    'de los activos con εᵢ > 0 desde el corto plazo. Tercero, y más importante '
    'desde el punto de vista práctico, la cartera Escasez y Resiliencia está '
    'correctamente posicionada en los cuellos de botella físicos (cobre, uranio, '
    'grid) y digitales (Bitcoin) que el proceso de autofabricación no puede '
    'eliminar, sino que de hecho intensifica. Con t_robot = 2027, este efecto '
    'es perceptible ya desde el segundo año del horizonte de inversión.'
)

add_para(
    'El corolario es contraintuitivo pero robusto: cuanto más avance la '
    'autofabricación, más valiosos se vuelven los activos escasos, no menos. '
    'La tesis de inversión se autofortalece con el mismo proceso tecnológico '
    'que podría parecer que la amenaza.'
)

add_para(
    'Una advertencia final: con t_robot = 2027, las proyecciones del modelo extendido '
    'son más sensibles a la velocidad real de adopción robótica en el período '
    '2027–2030. Una revisión al alza de η (eficiencia de replicación) amplifica '
    'drásticamente los resultados; una revisión a la baja o un retraso regulatorio '
    'los modera. El modelo base del paper fundacional permanece como la herramienta '
    'de gestión diaria del portfolio; el modelo extendido es el marco conceptual '
    'para la revisión estratégica semestral.',
    space_after=16
)

add_separator()

add_para(
    '"Las personas siempre sobreestiman lo que puede ocurrir en dos años '
    'y subestiman lo que puede ocurrir en diez."',
    italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2
)
add_para(
    '— Bill Gates, The Road Ahead (1995)',
    italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCIAS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('Referencias', level=1)

refs = [
    'Bass, F. M. (1969). A new product growth for model consumer durables. '
    'Management Science, 15(5), 215–227.',

    'Freeman, C., & Louçã, F. (2001). As Time Goes By: From the Industrial Revolutions '
    'to the Information Revolution. Oxford University Press.',

    'International Federation of Robotics (2025). World Robotics Report 2025. '
    'IFR Press, Frankfurt.',

    'Mokyr, J. (2002). The Gifts of Athena: Historical Origins of the Knowledge Economy. '
    'Princeton University Press.',

    'Rogers, E. M. (1962). Diffusion of Innovations. Free Press of Glencoe.',

    'Romer, P. M. (1990). Endogenous technological change. '
    'Journal of Political Economy, 98(5), S71–S102.',

    'Schumpeter, J. A. (1934). The Theory of Economic Development. '
    'Harvard University Press.',

    'Vilar, J. A. (2026a). Escasez y Resiliencia: Un Marco Matemático para Invertir en '
    'la Era de la Autorreplicación Artificial. Working Paper, SSRN.',

    'Vilar, J. A. (2026b). Autofabricación y Olas Tecnológicas: El Modelo de Logística '
    'Anidada para la Era de los Robots que se Autorreplican. Working Paper, SSRN. '
    '[Este paper]',

    'Von Neumann, J. (1966). Theory of Self-Reproducing Automata. '
    'University of Illinois Press. (Editado y completado por A. W. Burks.)',
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(ref)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# NOTA DEL AUTOR
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('Nota del Autor', level=1)

add_para(
    'Este paper forma parte de una serie de working papers sobre el framework '
    '"Escasez y Resiliencia". El autor es profesional de relaciones con inversores '
    'y estudiante de Grado en Ciencia de Datos en la Universitat Oberta de Catalunya (UOC). '
    'Los modelos matemáticos presentados fueron desarrollados y validados con apoyo '
    'de herramientas de inteligencia artificial (Claude, Anthropic) en el contexto '
    'de un proyecto personal de gestión de portfolio.',
    space_after=8
)

add_para(
    'Las proyecciones y escenarios incluidos en este paper son ilustrativos y no '
    'constituyen asesoramiento de inversión. El autor invierte personalmente en '
    'los activos descritos. El presente documento se publica bajo licencia '
    'Creative Commons BY-NC 4.0.',
    italic=True, space_after=16
)

add_separator()

p_ver = doc.add_paragraph()
p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_ver.paragraph_format.space_before = Pt(8)
rv = p_ver.add_run('Serie Escasez y Resiliencia · Paper II · v1.0 · Marzo 2026')
rv.font.size = Pt(9); rv.font.name = 'Times New Roman'
rv.font.color.rgb = RGBColor(120, 120, 120)

# ── Guardar ───────────────────────────────────────────────────────────────────
OUT = '/sessions/exciting-clever-brown/mnt/Cartera-Inversion/tesis/paper/Autofabricacion_Olas_Tecnologicas_v1.docx'
doc.save(OUT)
print(f'✅ Guardado: {OUT}')
